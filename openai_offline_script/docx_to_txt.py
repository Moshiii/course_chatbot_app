import os

import docx

def getText(filename):
    print(filename)

    doc = docx.Document(filename)
    # print(len(doc.paragraphs))
    fullText = []
    for para in doc.paragraphs:
        # print(len(para.text))
        fullText.append(para.text)
    return '\n'.join(fullText)


def read_all_docx_files_in_folder(folder_path):
    file_list = []
    for file in os.listdir(folder_path):
        if file.endswith(".docx"):
            file_list.append(folder_path + file)
    return file_list

path = "C:\@code\course_chatbot_app\content\\"
file_list = read_all_docx_files_in_folder(path)
for f in file_list:
    print(f)
    text = getText(f)
    print(len(text))
    with open(f.split('.')[0]
              +".txt", "w",encoding='utf-8') as text_file:
        text_file.write(text)

